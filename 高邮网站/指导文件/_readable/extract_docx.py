# -*- coding: utf-8 -*-
import docx
p = r"D:\新文科比赛\高邮网站\网站设计稿.docx"
d = docx.Document(p)
print("=== 段落 ===")
for para in d.paragraphs:
    t = para.text.strip()
    if t:
        print(t)
print("\n=== 表格 ===")
for ti, table in enumerate(d.tables):
    print(f"--- 表格 {ti+1} ---")
    for row in table.rows:
        cells = [c.text.strip().replace("\n", " / ") for c in row.cells]
        print(" | ".join(cells))
