from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ===== Title =====
title = doc.add_heading('', level=0)
run = title.add_run('技术路线')
run.font.size = Pt(22)
run.font.bold = True
run.font.color.rgb = RGBColor(0x17, 0x3a, 0x2e)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('高邮湖泊湿地多类型遗产数字活化平台')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x2d, 0x6a, 0x4f)
run.font.bold = True

doc.add_paragraph('')

# ===== 3-1 项目总体实施思路 =====
doc.add_heading('3-1 项目总体实施思路（介绍项目的总体思路、技术路线、开展实践的方法等，不超过500字）', level=1)

doc.add_paragraph(
    '本项目以高邮湖泊湿地多类型遗产为核心，依托 SQL Server、Java 及 GIS 技术，'
    '构建"调研梳理—技术开发—场景落地—客群适配"的完整实施路径。'
)

doc.add_paragraph(
    '首先，调研梳理与多模态数据转化。通过文献研究与实地走访，'
    '提取"生态链结构""传统农技""文化内涵"三大核心要素。'
    '将调研成果转化为多层次可开发数据：利用 SQL Server 构建完整数据库，'
    '存储 215 份有效问卷、10 人次深度访谈、800 余张调研照片与 11 处遗产点位坐标；'
    '将"稻鸭鱼蟹共作""里运河—高邮灌区闸坝体系""高邮民歌与咸鸭蛋技艺"等'
    '农技流程拆解为可交互的逻辑脚本与数据库字段，'
    '为跨平台开发提供数据支撑。我们将利用 SQL Server 构建出完整数据库，并将其公开供需要者使用。'
)

doc.add_paragraph(
    '其次，基于 MapLibre GL 与 ECharts 构建可视化展示核心架构。'
    '遵循"探索—深入—收获"流程进行开发，'
    '通过 Java 后端 API 的数据交互，实现经验值兑换实体衍生品、打卡分享获折扣等功能，'
    '增强用户粘性。通过 MapLibre GL 实现 3D 交互式高邮湖地图，11 处遗产点位带图文弹窗，'
    '并嵌入 CSS 动画系统，在关键节点弹出文化知识点解释。'
)

doc.add_paragraph(
    '最后，构建多端联动的农遗传承模式。'
    '利用 Unity 的跨平台发布能力，一键生成 iOS、Android 及 PC 客户端。'
    '针对文化传播场景，在 Unity 中内置知识问答事件系统；'
    '针对文旅体验场景，开发线下打卡功能，'
    '关联核心种植区坐标，实现线上线下联动。'
    '所有玩家数据通过 Java 后端统一管理，'
    '确保"线上互动+线下场景+知识传播"模式稳定运行，'
    '最终形成可持续农遗传承生态。'
)

doc.add_paragraph('')

# ===== 3-2 数据库与智能关联功能实现 =====
doc.add_heading('3-2 数据库与智能关联功能实现', level=1)

t1 = doc.add_table(rows=8, cols=2, style='Light Grid Accent 1')
t1.alignment = WD_TABLE_ALIGNMENT.CENTER
db_rows = [
    ('功能模块', '实现技术'),
    ('气候、水位、湿地面积等时序数据管理',
     'SQL Server + ETL工具（如Kettle）进行数据清洗与入库'),
    ('多维度数据检索与筛选',
     'Elasticsearch（全文检索引擎）/ SQL Server全文索引'),
    ('数据关联分析与预警模型（气候→水位耦合分析）',
     'Python（Pandas + Scikit-learn）进行回归分析与趋势建模'),
    ('后端数据接口与业务逻辑',
     'Java（Spring Boot）构建RESTful API'),
    ('前后端数据交互',
     'Axios / Fetch调用API，JSON格式传输'),
    ('地图与数据看板前端展示',
     'Vue.js / React + ECharts（数据可视化图表库）'),
    ('水位预警与观景推荐触发',
     '规则引擎（Drools / 硬编码阈值判断）+ 消息推送（WebSocket）'),
]
for i, (a, b) in enumerate(db_rows):
    t1.rows[i].cells[0].text = a
    t1.rows[i].cells[1].text = b
    for j in range(2):
        for p in t1.rows[i].cells[j].paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)
                if i == 0:
                    r.font.bold = True

doc.add_paragraph('')

# ===== 3-3 IP 形象与文创设计 =====
doc.add_heading('3-3 IP形象与文创设计', level=1)

t2 = doc.add_table(rows=5, cols=2, style='Light Grid Accent 1')
t2.alignment = WD_TABLE_ALIGNMENT.CENTER
ip_rows = [
    ('功能模块', '实现技术'),
    ('IP形象设计文件制作',
     'Adobe Illustrator / Photoshop（SVG矢量图形制作）'),
    ('IP形象动态展示与交互动效',
     'CSS动画 / Three.js（网页端3D展示） / Lottie（矢量动画）'),
    ('IP形象文创产品延展设计',
     'Adobe InDesign / Photoshop排版 + 印刷供应链对接'),
]
for i, (a, b) in enumerate(ip_rows):
    t2.rows[i].cells[0].text = a
    t2.rows[i].cells[1].text = b
    for j in range(2):
        for p in t2.rows[i].cells[j].paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)
                if i == 0:
                    r.font.bold = True

doc.add_paragraph('')

# ===== 3-4 交互式叙事游戏 =====
doc.add_heading('3-4 交互式叙事游戏', level=1)

t3 = doc.add_table(rows=6, cols=2, style='Light Grid Accent 1')
t3.alignment = WD_TABLE_ALIGNMENT.CENTER
game_rows = [
    ('功能模块', '实现技术'),
    ('游戏引擎选型（3D场景与跨平台能力）',
     'Unity（支持iOS/Android/PC/Web跨平台发布）'),
    ('2D场景与叙事界面',
     'Unity 2D / 微信小游戏框架 / HTML5 Canvas'),
    ('剧情脚本与对话系统',
     'Unity Timeline + Cinemachine / 自研可视化剧情编辑器'),
    ('经验值、成就、背包等玩家系统',
     'Java Spring Boot后端 + MySQL数据持久化'),
    ('文化知识点弹窗与百科系统',
     'Unity UI / 网页端overlay浮层 + 后端CMS内容管理'),
]
for i, (a, b) in enumerate(game_rows):
    t3.rows[i].cells[0].text = a
    t3.rows[i].cells[1].text = b
    for j in range(2):
        for p in t3.rows[i].cells[j].paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)
                if i == 0:
                    r.font.bold = True

doc.add_paragraph('')

# ===== 3-5 湿地实景剧本杀 =====
doc.add_heading('3-5 湿地实景剧本杀', level=1)

t4 = doc.add_table(rows=6, cols=2, style='Light Grid Accent 1')
t4.alignment = WD_TABLE_ALIGNMENT.CENTER
script_rows = [
    ('功能模块', '实现技术'),
    ('剧本杀App/小程序载体',
     '微信小程序（微信开发者工具）/ 独立App（Flutter / React Native）'),
    ('LBS定位与打卡签到',
     '微信小程序wx.getLocation / 高德定位SDK，结合坐标围栏算法判断用户是否在点位附近'),
    ('AR线索识别',
     '微信小程序AR插件 / 百度AI识别 / EasyAR'),
    ('剧情解锁与任务分发',
     '后端根据用户GPS坐标 + 任务状态动态下发剧情节点'),
    ('图文/音频/视频线索加载',
     '小程序原生媒体组件 + 云存储（OSS）存储线索素材'),
]
for i, (a, b) in enumerate(script_rows):
    t4.rows[i].cells[0].text = a
    t4.rows[i].cells[1].text = b
    for j in range(2):
        for p in t4.rows[i].cells[j].paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)
                if i == 0:
                    r.font.bold = True

doc.add_paragraph('')

# ===== 3-6 IP 门户网站 =====
doc.add_heading('3-6 IP门户网站', level=1)

t5 = doc.add_table(rows=6, cols=2, style='Light Grid Accent 1')
t5.alignment = WD_TABLE_ALIGNMENT.CENTER
web_rows = [
    ('功能模块', '实现技术'),
    ('网站前端展示',
     'HTML5 + CSS3 + JavaScript / Vue.js / React'),
    ('IP形象动态展示与交互动效',
     'CSS动画 / Three.js（网页端3D展示） / Lottie（矢量动画）'),
    ('多类型遗产资源图谱展示',
     '嵌入Leaflet/高德地图组件，点击标注点弹出图文介绍浮层'),
    ('网站后端内容管理',
     '轻量级CMS / Java Spring Boot + Thymeleaf'),
    ('游戏与剧本杀入口跳转',
     'HTTPS链接跳转 / OAuth单点登录（实现统一账号体系）'),
]
for i, (a, b) in enumerate(web_rows):
    t5.rows[i].cells[0].text = a
    t5.rows[i].cells[1].text = b
    for j in range(2):
        for p in t5.rows[i].cells[j].paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)
                if i == 0:
                    r.font.bold = True

# Save
output_path = r'D:\新文科比赛\技术路线_v2.docx'
doc.save(output_path)
print('Word document saved to:', output_path)
print('File size:', os.path.getsize(output_path), 'bytes')
